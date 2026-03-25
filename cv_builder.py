import streamlit as st
import openai
from datetime import datetime
import json
import docx
from docx import Document
from docx.shared import Inches
import pandas as pd
from io import BytesIO

# Configure OpenAI API (replace with your key)
openai.api_key = st.secrets.get("OPENAI_API_KEY", "your-api-key-here")

class CVBuilderAgent:
    def __init__(self):
        self.sections = {
            'personal_info': {},
            'summary': '',
            'experience': [],
            'education': [],
            'skills': [],
            'projects': [],
            'certifications': []
        }
    
    def generate_section(self, section_type, user_input, context=""):
        """Generate content for specific CV section using AI"""
        prompts = {
            'summary': f"""Write a professional 4-6 sentence career summary for someone with this background: {user_input}. 
            Make it concise, impactful, and tailored for {context}. Use action verbs and quantify achievements where possible.""",
            
            'experience': f"""Convert this job experience into 4-6 bullet points with strong action verbs: {user_input}. 
            Use the STAR method (Situation, Task, Action, Result). Quantify achievements. Format as bullet points.""",
            
            'skills': f"""Based on this experience: {user_input}, suggest 8-12 relevant technical and soft skills. 
            Categorize them (Technical, Soft Skills, Tools). Return as a comma-separated list.""",
            
            'projects': f"""Transform this project description into 3-5 impressive bullet points: {user_input}. 
            Highlight technologies used, your role, challenges solved, and results achieved.""",
            
            'education': f"""Format this education information professionally: {user_input}. 
            Include relevant coursework, GPA (if >3.5), honors, and achievements."""
        }
        
        prompt = prompts.get(section_type, f"Generate professional CV content for {section_type}: {user_input}")
        
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a professional resume writer with 20+ years experience. Write concise, impactful content optimized for ATS systems."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=500,
            temperature=0.7
        )
        
        return response.choices[0].message.content.strip()
    
    def optimize_keywords(self, job_description, cv_content):
        """Extract keywords from job description and suggest improvements"""
        prompt = f"""Analyze this job description: {job_description}
        
        Current CV content: {cv_content}
        
        Extract 10-15 key skills, technologies, and qualifications. Suggest how to incorporate them naturally into the CV. Return as:
        1. Missing keywords list
        2. Suggested improvements (3-5 specific changes)"""
        
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an ATS optimization expert. Identify missing keywords and suggest natural incorporation."},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content.strip()
    
    def generate_cv(self):
        """Generate complete CV content"""
        cv_content = f"""
# {self.sections['personal_info'].get('name', 'Your Name')}

{self.sections['personal_info'].get('title', '')} | {self.sections['personal_info'].get('email', '')} | {self.sections['personal_info'].get('phone', '')} | {self.sections['personal_info'].get('linkedin', '')} | {self.sections['personal_info'].get('location', '')}

## Professional Summary
{self.sections['summary']}

## Professional Experience
"""
        
        for exp in self.sections['experience']:
            cv_content += f"\n### {exp.get('title', '')} | {exp.get('company', '')} | {exp.get('dates', '')}\n{exp.get('description', '')}\n"
        
        cv_content += "\n## Education\n"
        for edu in self.sections['education']:
            cv_content += f"### {edu.get('degree', '')}, {edu.get('field', '')}\n{edu.get('school', '')} | {edu.get('dates', '')}\n{edu.get('details', '')}\n"
        
        cv_content += "\n## Skills\n" + ", ".join(self.sections['skills'])
        
        return cv_content

def create_word_doc(cv_content, personal_info):
    """Create professional Word document"""
    doc = Document()
    
    # Header
    header = doc.add_heading(personal_info.get('name', 'Your Name'), 0)
    header.alignment = 1  # Center
    
    subtitle = doc.add_paragraph()
    subtitle.add_run(personal_info.get('title', '')).bold = True
    subtitle.add_run(f" | {personal_info.get('email', '')} | {personal_info.get('phone', '')} | {personal_info.get('linkedin', '')}").italic = True
    
    # Sections
    sections = cv_content.split('## ')[1:]
    for section in sections:
        title, content = section.split('\n', 1)
        doc.add_heading(title.strip(), level=1)
        doc.add_paragraph(content.strip())
    
    # Save to BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Streamlit UI
def main():
    st.set_page_config(page_title="AI CV Builder", layout="wide")
    st.title("🚀 AI-Powered CV Builder")
    
    agent = CVBuilderAgent()
    
    # Sidebar for personal info
    with st.sidebar:
        st.header("Personal Information")
        personal_info = {
            'name': st.text_input("Full Name", "John Doe"),
            'title': st.text_input("Job Title", "Software Engineer"),
            'email': st.text_input("Email", "john.doe@email.com"),
            'phone': st.text_input("Phone", "+1-234-567-8900"),
            'linkedin': st.text_input("LinkedIn", "linkedin.com/in/johndoe"),
            'location': st.text_input("Location", "San Francisco, CA")
        }
        agent.sections['personal_info'] = personal_info
    
    tab1, tab2, tab3, tab4 = st.tabs(["📝 Build CV", "⚡ AI Generate", "🎯 Job Match", "📄 Download"])
    
    with tab1:
        st.header("Step 1: Enter Your Information")
        
        # Experience
        st.subheader("Work Experience")
        for i in range(3):  # Up to 3 experiences
            with st.expander(f"Experience {i+1}"):
                col1, col2 = st.columns(2)
                with col1:
                    title = st.text_input(f"Title {i+1}", key=f"title_{i}")
                    company = st.text_input(f"Company {i+1}", key=f"company_{i}")
                with col2:
                    dates = st.text_input(f"Dates {i+1}", key=f"dates_{i}")
                
                desc = st.text_area(f"Description {i+1}", key=f"desc_{i}")
                if st.button(f"✨ AI Enhance Experience {i+1}", key=f"enhance_exp_{i}"):
                    if desc:
                        enhanced = agent.generate_section('experience', desc)
                        st.text_area("AI Enhanced:", value=enhanced, key=f"enhanced_exp_{i}")
        
        # Education
        st.subheader("Education")
        for i in range(2):
            with st.expander(f"Education {i+1}"):
                degree = st.text_input(f"Degree {i+1}", key=f"degree_{i}")
                school = st.text_input(f"School {i+1}", key=f"school_{i}")
                dates = st.text_input(f"Dates {i+1}", key=f"edu_dates_{i}")
                details = st.text_area(f"Details {i+1}", key=f"edu_details_{i}")
    
    with tab2:
        st.header("AI Content Generation")
        section_type = st.selectbox("Select section to generate", 
                                  ["Professional Summary", "Skills", "Experience Bullet", "Project Description"])
        
        user_input = st.text_area("Your raw input", "Tell me about your experience...")
        context = st.text_input("Target job title", "Software Engineer")
        
        if st.button("✨ Generate Content"):
            with st.spinner("Generating with AI..."):
                generated = agent.generate_section(section_type.lower().replace(' ', '_'), 
                                                 user_input, context)
                st.markdown("### Generated Content")
                st.write(generated)
                st.download_button("Copy to clipboard", generated)
    
    with tab3:
        st.header("🎯 Job Description Matcher")
        job_desc = st.text_area("Paste Job Description", height=200)
        current_cv = st.text_area("Your Current CV", height=200)
        
        if st.button("🔍 Analyze & Optimize"):
            if job_desc and current_cv:
                with st.spinner("Analyzing..."):
                    optimization = agent.optimize_keywords(job_desc, current_cv)
                    st.markdown("### Optimization Recommendations")
                    st.write(optimization)
    
    with tab4:
        st.header("📄 Download Your CV")
        if st.button("Generate Full CV Preview"):
            # Simple preview
            cv_preview = agent.generate_cv()
            st.markdown(cv_preview)
            
            # Download Word doc
            word_file = create_word_doc(cv_preview, personal_info)
            st.download_button(
                "⬇️ Download Word CV",
                word_file,
                f"CV_{personal_info['name'].replace(' ', '_')}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # Download Markdown
            st.download_button(
                "⬇️ Download Markdown",
                cv_preview,
                f"CV_{personal_info['name'].replace(' ', '_')}.md",
                "text/markdown"
            )

if __name__ == "__main__":
    main()
